from datetime import date, timedelta, datetime
from pages.calculadora_prazo import transforma_datas, dia_util
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def gerar_relatorio_prazo(data_inicial, data_final, lista_inicio):
    contador_par = 0
    datas_do_prazo = []
    dia_do_prazo = []
    eventos = []
    lista_relatorio = [['Data', 'Dia do prazo', 'Evento']]
    data_inicial = transforma_datas.str_para_date(data_inicial)
    data_final = transforma_datas.str_para_date(data_final)
    for data in lista_inicio:
        data_str = transforma_datas.date_para_str(data)
        lista_relatorio.append([data_str, "-", "Prazo não iniciado"])
    while data_inicial <= data_final:
        data_str = data_inicial
        if dia_util.verificador_dia_util(data_str):
            data_str = transforma_datas.date_para_str(data_str)
            contador_par += 1
            datas_do_prazo.append(data_str)
            dia_do_prazo.append(f'{str(contador_par)}')
            evento = ''
            eventos.append(evento)
            lista_relatorio.append([datas_do_prazo[-1], dia_do_prazo[-1], eventos[-1]])
        else:
            evento = dia_util.verificar_fds_feriado_relatorio(data_str)
            data_str = transforma_datas.date_para_str(data_str)
            datas_do_prazo.append(data_str)
            dia_do_prazo.append('-')
            eventos.append(evento)
            lista_relatorio.append([datas_do_prazo[-1], dia_do_prazo[-1], eventos[-1]])
        data_inicial += timedelta(days=1)
    gerar_docx(lista_relatorio)


def gerar_docx(lista_relatorio):
    relatorio = Document()
    titulo = relatorio.add_heading('Relatório de prazos', 1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabela = relatorio.add_table(rows=1, cols=3)
    tabela.style = 'Medium Shading 1 Accent 1'

    # Cabeçalho centralizado e em negrito
    hdr_cells = tabela.rows[0].cells
    for i, c in enumerate(lista_relatorio[0]):
        run = hdr_cells[i].paragraphs[0].add_run(c)
        run.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Dados centralizados e sem negrito
    for linha in lista_relatorio[1:]:
        linha_celulas = tabela.add_row().cells
        for i, c in enumerate(linha):
            run = linha_celulas[i].paragraphs[0].add_run(c)
            run.bold = False  # Garante que não fique em negrito
            linha_celulas[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    relatorio.save('Relatório.docx')