from docx import Document
import os
from pathlib import Path

def gerar_ci(origem,
             destino,
             assunto,
             data_ci_str,
             numero_ci_estilizado,
             servidor,
             cargo,
             matricula,
             texto):

    parametros = {
    '@De_origem': origem,
    '@Para_destino': destino,
    '@Assunto': assunto,
    '@Data': data_ci_str,
    '@NúmeroCI': numero_ci_estilizado,
    '@Nome': servidor,
    '@Cargo': cargo,
    '@Matrícula': matricula,
    '@Texto': texto,
                }    

    current_dir = Path(__file__).parent
    modelo_path = current_dir / 'CI_Modelo.docx'

    certidao = Document(modelo_path)

    def substituir_parametros(paragrafo):
        """Função auxiliar para substituir parâmetros em um parágrafo"""
        for parametro, valor in parametros.items():
            if parametro in paragrafo.text:
                for run in paragrafo.runs:
                    if parametro in run.text:
                        run.text = run.text.replace(parametro, valor)    

    for tabela in certidao.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    substituir_parametros(paragrafo)
                
                # **Verifica também se há tabelas dentro de células (aninhadas)**
                for tabela_aninhada in celula.tables:
                    for linha_aninhada in tabela_aninhada.rows:
                        for celula_aninhada in linha_aninhada.cells:
                            for paragrafo_aninhado in celula_aninhada.paragraphs:
                                substituir_parametros(paragrafo_aninhado)

    certidao.save(f'CI - {numero_ci_estilizado} - {destino} - {assunto}.docx')
    return (f'CI - {numero_ci_estilizado} - {destino} - {assunto}.docx')
    
  
        
