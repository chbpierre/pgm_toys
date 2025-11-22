from docx import Document
import os
from pathlib import Path



def gerar_certidão(opcao_tipo_certidao, procurador, oab, processo, data_certidão, parte_contraria):

    parametros = {
    '@NúmeroProcesso': processo,
    '@NomeAutor': parte_contraria,
    '@Data': data_certidão,
    '@NomeProcurador': procurador,
    '@OAB': oab,
                }    

    if opcao_tipo_certidao == 'ATS Hora extra':
        
        current_dir = Path(__file__).parent
        modelo_path = current_dir / 'Certidão_ATS.docx'
        certidao = Document(modelo_path)

        for paragrafo in certidao.paragraphs:
            for parametro, valor in parametros.items():
                if parametro in paragrafo.text:
                    for run in paragrafo.runs:
                        if parametro in run.text:
                            run.text = run.text.replace(parametro, valor)

        certidao.save(f'Certidão ATS - {parte_contraria}.docx')
        return (f'Certidão ATS - {parte_contraria}.docx')
    
    elif opcao_tipo_certidao == "PCCV Geral - Recurso Inominado":
       
        current_dir = Path(__file__).parent
        modelo_path = current_dir / 'Certidão_RI_PCCV_Geral.docx'
        certidao = Document(modelo_path)

        for paragrafo in certidao.paragraphs:
            for parametro, valor in parametros.items():
                if parametro in paragrafo.text:
                    for run in paragrafo.runs:
                        if parametro in run.text:
                            run.text = run.text.replace(parametro, valor)

        certidao.save(f'Certidão PCCV Geral - RI - {parte_contraria}.docx')
        return (f'Certidão PCCV Geral - RI - {parte_contraria}.docx')
    
    elif opcao_tipo_certidao == 'PCCV Magistério - Recurso Inominado':

        current_dir = Path(__file__).parent
        modelo_path = current_dir / 'Certidão_RI_PCCV_Magistério.docx'
        certidao = Document(modelo_path)

        for paragrafo in certidao.paragraphs:
            for parametro, valor in parametros.items():
                if parametro in paragrafo.text:
                    for run in paragrafo.runs:
                        if parametro in run.text:
                            run.text = run.text.replace(parametro, valor)

        certidao.save(f'Certidão PCCV Magistério - RI - {parte_contraria}.docx')
        return (f'Certidão PCCV Magistério - RI - {parte_contraria}.docx')
    
    elif opcao_tipo_certidao == 'PCCV Geral - Apelação':

        current_dir = Path(__file__).parent
        modelo_path = current_dir / 'Certidão_Apelação_PCCV_Geral.docx'
        certidao = Document(modelo_path)

        for paragrafo in certidao.paragraphs:
            for parametro, valor in parametros.items():
                if parametro in paragrafo.text:
                    for run in paragrafo.runs:
                        if parametro in run.text:
                            run.text = run.text.replace(parametro, valor)

        certidao.save(f'Certidão PCCV Geral - Apelação - {parte_contraria}.docx')
        return (f'Certidão PCCV Geral - Apelação - {parte_contraria}.docx')
    
        
